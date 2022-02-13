import docx
from docx.enum.text import WD_COLOR_INDEX
from utilities import get_transcripts, TranscriptType

def read_results(path):
    results = list()
    entities = list()

    with open(path) as file:
        line = file.readline()
        while line != '':
            if line == 'batch_end\n':
                results.append(entities)
                line = file.readline()
                entities = list()
                continue
            entities.append(eval(line))
            line = file.readline()

    return results

def write_to_doc(transcript, results):
    for batch, entities in zip(transcript, results):
        para = doc.add_paragraph()
        prev_idx = 0

        for entity in entities:
            start_idx = entity[0][0]
            end_idx = entity[0][1]
            entity_class = entity[1]

            color = WD_COLOR_INDEX.YELLOW

            if entity_class == 'PER':
                color = WD_COLOR_INDEX.RED
            elif entity_class == 'LOC':
                color = WD_COLOR_INDEX.GREEN
            elif entity_class == 'ORG':
                color = WD_COLOR_INDEX.TURQUOISE

            para.add_run(batch[prev_idx:start_idx])
            para.add_run(batch[start_idx:end_idx]).font.highlight_color = color
            prev_idx = end_idx


################################################################
# Main Function

if __name__ == '__main__':
    # Read the transcript in
    directory = "../transcripts/ingested"
    transcript = get_transcripts(TranscriptType.TEST, directory)

    doc = docx.Document()
    para = doc.add_paragraph('LEGEND:\n')
    para.add_run('PERSON\n').font.highlight_color = WD_COLOR_INDEX.RED
    para.add_run('LOCATION\n').font.highlight_color = WD_COLOR_INDEX.GREEN
    para.add_run('ORGANISATION\n').font.highlight_color = WD_COLOR_INDEX.TURQUOISE
    para.add_run('OTHER\n').font.highlight_color = WD_COLOR_INDEX.YELLOW

    # Visualise GATE results
    results = read_results('./outputs/gate_results.txt')
    doc.add_heading('GATE', 0)
    write_to_doc(transcript, results)
    
    # Visualise Flair results
    results = read_results('./outputs/flair_results.txt')
    doc.add_heading('Flair', 0)
    write_to_doc(transcript, results)

    # Visualise spaCy results
    results = read_results('./outputs/spacy_results.txt')
    doc.add_heading('spaCy', 0)
    write_to_doc(transcript, results)

    # Visualise Stanford results
    results = read_results('./outputs/stanford_results.txt')
    doc.add_heading('Stanford', 0)
    write_to_doc(transcript, results)

    # Visualise DeepPavlov results
    # TODO

    # Save the document
    doc.save('./outputs/highlighted_entities.docx')